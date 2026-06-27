"""
Word 文档读取器 (.docx)

读取 .docx 文件中的文字, 包括段落和表格

后续 C++ pybind11 对应:
- C++ 端使用 miniz + pugixml 解析 docx 的 ZIP/XML 结构
"""

from .base import BaseReader

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DocxReader(BaseReader):
    """读取 .docx 文件中的文字内容"""

    SUPPORTED = {".docx"}

    @property
    def supported_extensions(self) -> list:
        return sorted(self.SUPPORTED)

    def read(self, file_path: str) -> str:
        """
        提取 .docx 文件中的所有文字

        包括:
        - 正文段落
        - 表格中的文字
        - 页眉页脚 (如果存在)

        Args:
            file_path: .docx 文件路径

        Returns:
            str: 提取的全部文字, 段落间换行
        """
        if not HAS_DOCX:
            raise ImportError(
                "python-docx 未安装, 请执行: pip install python-docx"
            )

        doc = Document(file_path)
        parts = []

        # 1. 正文段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # 2. 表格内容
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    parts.append(" | ".join(row_texts))

        # 3. 页眉页脚 (如有)
        for section in doc.sections:
            # 页眉
            header = section.header
            if header and not header.is_linked_to_previous:
                for para in header.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(f"[页眉] {text}")
            # 页脚
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                for para in footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(f"[页脚] {text}")

        return "\n".join(parts)
