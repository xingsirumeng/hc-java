"""
Smoke tests for local file processing and the HTTP service factory.
"""

import importlib.util
import os
import sys
import tempfile


sys.path.insert(0, os.path.join(os.path.dirname(__file__), "src"))

from hc_re_ai import JobProcessor, detect_file_type


class SkipTest(Exception):
    pass


def has_module(module_name: str) -> bool:
    return importlib.util.find_spec(module_name) is not None


def test_detect_file_type():
    print("=" * 50)
    print("[TEST] detect_file_type")

    assert detect_file_type("test.txt") == "txt"
    assert detect_file_type("test.docx") == "docx"
    assert detect_file_type("test.png") == "image"
    assert detect_file_type("test.jpg") == "image"
    assert detect_file_type("test.jpeg") == "image"
    assert detect_file_type("test.PNG") == "image"
    assert detect_file_type("test.xyz") == "unknown"

    print("  PASS")


def test_txt():
    print("=" * 50)
    print("[TEST] txt")

    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".txt", delete=False, encoding="utf-8"
    ) as temp_file:
        temp_file.write("Hello World\nThis is a smoke test\nLine 3")
        tmp_path = temp_file.name

    try:
        processor = JobProcessor()
        result = processor.process(tmp_path)

        assert result.success, result.error
        assert "Hello World" in result.text
        assert "smoke test" in result.text
        assert result.file_type == "txt"
        print("  PASS")
    finally:
        os.unlink(tmp_path)


def test_docx():
    print("=" * 50)
    print("[TEST] docx")

    if not has_module("docx"):
        raise SkipTest("python-docx is not installed")

    from docx import Document

    tmp_path = os.path.join(tempfile.gettempdir(), "test_hc_re_ai.docx")
    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph Hello DOCX")
    doc.save(tmp_path)

    try:
        processor = JobProcessor()
        result = processor.process(tmp_path)

        assert result.success, result.error
        assert "First paragraph" in result.text
        assert "Hello DOCX" in result.text
        assert result.file_type == "docx"
        print("  PASS")
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def test_file_not_found():
    print("=" * 50)
    print("[TEST] file_not_found")

    processor = JobProcessor()
    result = processor.process("D:/nonexistent_file_xyz123.txt")

    assert not result.success
    assert "file does not exist" in result.error
    print("  PASS")


def test_unsupported_type():
    print("=" * 50)
    print("[TEST] unsupported_type")

    with tempfile.NamedTemporaryFile(mode="w", suffix=".xyz", delete=False) as temp_file:
        temp_file.write("test")
        tmp_path = temp_file.name

    try:
        processor = JobProcessor()
        result = processor.process(tmp_path)

        assert result.success, result.error
        assert result.file_type == "txt (fallback)"
        print("  PASS")
    finally:
        os.unlink(tmp_path)


def test_batch():
    print("=" * 50)
    print("[TEST] batch")

    file_paths = []
    for index in range(3):
        temp_file = tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False, encoding="utf-8"
        )
        temp_file.write("Batch file %d\nContent line" % (index + 1))
        temp_file.close()
        file_paths.append(temp_file.name)

    try:
        processor = JobProcessor()
        results = processor.process_batch(file_paths)

        assert len(results) == 3
        for result in results:
            assert result.success, result.error
            assert result.file_type == "txt"
            assert "Batch file" in result.text
        print("  PASS")
    finally:
        for path in file_paths:
            os.unlink(path)


def test_image():
    print("=" * 50)
    print("[TEST] image")

    if not has_module("easyocr"):
        raise SkipTest("easyocr is not installed")

    from PIL import Image, ImageDraw

    tmp_path = os.path.join(tempfile.gettempdir(), "test_hc_re_ai_ocr.png")
    image = Image.new("RGB", (500, 160), color="white")
    draw = ImageDraw.Draw(image)
    draw.text((20, 50), "Hello OCR Test 123", fill="black")
    image.save(tmp_path)

    try:
        processor = JobProcessor(
            ocr_lang_list=["en"],
            ocr_gpu=False,
            ocr_download_enabled=False,
        )
        result = processor.process(tmp_path)

        if not result.success:
            raise SkipTest("EasyOCR runtime/model is not ready: %s" % result.error)

        assert result.file_type == "image"
        assert len(result.text.strip()) > 0
        print("  PASS")
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def test_http_app_factory():
    print("=" * 50)
    print("[TEST] http_app_factory")

    if not has_module("fastapi"):
        raise SkipTest("fastapi is not installed")

    from hc_re_ai.http_service import create_app

    app = create_app(JobProcessor())
    routes = {route.path for route in app.routes}
    assert "/health" in routes
    assert "/api/v1/process/path" in routes
    assert "/api/v1/process/upload" in routes
    assert "/api/v1/process/batch" in routes
    print("  PASS")


def main():
    print("\n" + "=" * 50)
    print("hc_re_ai smoke tests")
    print("=" * 50 + "\n")

    tests = [
        test_detect_file_type,
        test_txt,
        test_docx,
        test_file_not_found,
        test_unsupported_type,
        test_batch,
        test_image,
        test_http_app_factory,
    ]

    passed = 0
    failed = 0
    skipped = 0

    for test_fn in tests:
        try:
            test_fn()
            passed += 1
        except SkipTest as exc:
            skipped += 1
            print("  SKIP: %s" % exc)
        except AssertionError as exc:
            failed += 1
            print("  FAIL: %s" % exc)
        except Exception as exc:
            failed += 1
            print("  ERROR: %s" % exc)

    print("\n" + "=" * 50)
    print("done: %d passed, %d failed, %d skipped" % (passed, failed, skipped))
    print("=" * 50)

    return failed == 0


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
